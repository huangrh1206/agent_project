import os
from docx import Document
from core.utils import log_markdown

# 约定属性
tool_name = "WordParser"
tool_desc = "读取并解析本地 .docx Word 文档的内容。参数为文件的绝对路径或相对路径。"

def tool_func(file_path: str) -> str:
    """
    解析 Word 文档并返回纯文本内容
    """
    log_markdown(f"📄 正在解析 Word 文档: {file_path}")
    
    if not os.path.exists(file_path):
        return f"错误：找不到文件 {file_path}，请确认路径是否正确。"
    
    if not file_path.lower().endswith('.docx'):
        return "错误：目前仅支持 .docx 格式的 Word 文档。"

    try:
        doc = Document(file_path)
        full_text = []
        
        # 1. 提取段落文字
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # 2. 提取表格文字 (可选)
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    full_text.append(" | ".join(row_data))

        result = "\n".join(full_text)
        
        if not result.strip():
            return "文档内容为空。"
            
        # 限制返回长度，防止撑爆 LLM 的上下文窗口
        return f"--- 文档内容开始 ---\n{result[:20000]}\n--- 文档内容结束 ---"

    except Exception as e:
        return f"解析 Word 时发生错误: {str(e)}"
